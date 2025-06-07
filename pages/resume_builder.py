import streamlit as st
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
import tempfile
import os
from dotenv import load_dotenv

from fpdf import FPDF

pdf = FPDF()

import os
print("Current working directory:", os.getcwd())

import unicodedata

def clean_unicode(text):
    return unicodedata.normalize('NFKD', text).encode('latin-1', 'ignore').decode('latin-1')
pdf.set_font("Arial", size=12)

# ========== Setup ==========
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

st.set_page_config(page_title="AI Resume Builder", layout="centered")
st.title("📝 AI Resume Builder")
st.markdown("Generate an ATS-friendly resume in seconds using AI.")

choice = st.radio("Start from:", ["📄 Upload Existing Resume", "🆕 Create from Scratch"])
job_title = st.text_input("Enter Target Job Title (e.g., Data Analyst)")
resume_content = ""

# ========== Extract Text ==========

def extract_text_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_pdf(file):
    import fitz  # PyMuPDF
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

# ========== Resume Input ==========

if choice == "📄 Upload Existing Resume":
    uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
    if uploaded_file and job_title:
        if uploaded_file.name.endswith(".pdf"):
            resume_content = extract_text_pdf(uploaded_file)
        else:
            resume_content = extract_text_docx(uploaded_file)

elif choice == "🆕 Create from Scratch":
    name = st.text_input("Your Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    summary = st.text_area("Professional Summary")
    skills = st.text_area("Skills (comma-separated)")
    experience = st.text_area("Experience (Job Title, Company, Duration, Description)")
    education = st.text_area("Education")
    projects = st.text_area("Projects (Optional)")

    if job_title:
        resume_content = f"""
        Name: {name}
        Email: {email}
        Phone: {phone}
        Summary: {summary}
        Skills: {skills}
        Experience: {experience}
        Education: {education}
        Projects: {projects}
        """

# ========== Prompt & AI Output ==========

prompt_template = f"""
You are a top-tier professional resume writer with deep expertise in ATS optimization, industry trends, and crafting highly effective resumes.

Rewrite and enhance the resume below for the job role: {job_title}.

Instructions:
- Use a clean, plain text format suitable for ATS (no tables, graphics, or complex formatting).
- Structure the resume with these sections:

1. Contact Information  
2. Professional Summary: 3-4 sentences tailored to the {job_title} role, highlighting key qualifications and career goals.  
3. Core Competencies: Group skills by categories, for example:  
   Programming Languages: Python, SQL, Java  
   Data Analysis Tools: Excel, Power BI, Tableau  
   Soft Skills: Communication, Problem Solving  
4. Professional Experience:  
   [Company Name], [Location]  
   [Job Title] | [Employment Dates]  
   • Use bullet points starting with strong action verbs.  
   • Emphasize achievements with quantifiable results where possible.  
   • Focus on responsibilities relevant to the {job_title} role.  
5. Projects:  
   Format each project as:  
   Project Name | Link  
   Brief 2-3 line description highlighting the scope, tools used, and outcome/results.  
6. Education:  
   [Degree], [Major]  
   [University], [Location]  
   [Graduation Date or Expected Graduation]  
7. Additional Sections (optional): Certifications, Professional Development, Awards, Languages, Volunteer Work

- Naturally incorporate relevant keywords from the job role.  
- Use concise, clear, and action-oriented language.  
- Ensure consistent formatting throughout.  
- Keep the resume focused and no longer than two plain text pages.  

Resume content to rewrite:  
------------------------  
{resume_content}  
------------------------

Return only the enhanced resume in plain text format. No explanations or extra commentary.
"""


def clean_unicode(text):
    return text.encode('latin-1', 'replace').decode('latin-1')

# ========== Generate Resume ==========
if resume_content.strip() != "" and st.button("✨ Enhance My Resume"):
    with st.spinner("Generating your optimized resume..."):
        prompt = prompt_template.format(job_title=job_title, resume_content=resume_content)
        response = model.generate_content(prompt)
        result = response.text

        st.success("Resume Generated Successfully!")
        st.text_area("📋 Enhanced Resume", result, height=400)

        # Save DOCX
        docx_path = os.path.join(tempfile.gettempdir(), "enhanced_resume.docx")
        doc = Document()
        for line in result.split('\n'):
            doc.add_paragraph(line.strip())
        doc.save(docx_path)

        def safe_multicell(pdf, text, width=None, height=10):
            text = text.strip()
            if not text:
                pdf.multi_cell(width or pdf.w - 2 * pdf.l_margin, height, "")
                return

    # Normalize and strip unsupported characters
            text = unicodedata.normalize('NFKD', text).encode('latin-1', 'ignore').decode('latin-1')

    # Break long unspaced lines (e.g., URLs or malformed chunks)
            if " " not in text and len(text) > 60:
                text = " ".join(text[i:i+50] for i in range(0, len(text), 50))

    # Try writing, fallback on placeholder if still erroring
            try:
                pdf.multi_cell(width or pdf.w - 2 * pdf.l_margin, height, text)
            except Exception:
                pdf.multi_cell(width or pdf.w - 2 * pdf.l_margin, height, "[Unrenderable Text]")



